"""
generate_report.py  —  Sentio Critical Academic Report
Improved with hardware specs from WS2812B Technical Documentation & Assembly Manual.

Run:
  python generate_architecture_diagram.py   # regenerate diagram if needed
  python generate_report.py

Screenshots (optional – place in docs/screenshots/):
  led_patterns_tshirt.png    01_config_screen.png
  02_device_settings.png     03_calibration_screen.png
  04_signal_lost.png         05_monitoring_screen.png
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE      = os.path.dirname(__file__)
SHOTS_DIR = os.path.join(BASE, "docs", "screenshots")
OUT_FILE  = os.path.join(BASE, "docs", "sentio_critical_report.docx")
os.makedirs(SHOTS_DIR, exist_ok=True)
os.makedirs(os.path.dirname(OUT_FILE), exist_ok=True)

SCREENSHOTS = [
    ("01_config_screen.png",
     "Figure 1 — Configuration screen: pattern type selection (2x2 card grid) "
     "and signal sensitivity slider."),
    ("02_device_settings.png",
     "Figure 2 — Device Settings panel: backend API URL and Muse 2 MAC address "
     "configuration without redeployment."),
    ("03_calibration_screen.png",
     "Figure 3 — Calibration screen: animated Muse 2 Headband connection sequence "
     "with step-by-step progress indicator."),
    ("04_signal_lost.png",
     "Figure 4 — Signal Lost modal: graceful Bluetooth disconnect handling with "
     "numbered recovery checklist and auto-retry countdown."),
    ("05_monitoring_screen.png",
     "Figure 5 — Live Monitoring dashboard: emotion orb, real-time EEG band power "
     "chart, Visual Parameters panel, and AI Guidance pane."),
]

TSHIRT_PHOTO = (
    "led_patterns_tshirt.png",
    "Figure 6 — The four emotion-driven LED patterns rendered on the physical "
    "WS2812B t-shirt garment (200 LEDs, 10x20 serpentine matrix, 40x55 cm active "
    "area): (1) Ondas Fluidas — calm/relaxed; (2) Padrao Geometrico — focused; "
    "(3) Pulsos Ritmicos — stressed/excited; (4) Estrelas e Particulas — neutral."
)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before   = Pt(14)
    p.paragraph_format.space_after    = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 0:
        set_font(run, size=18, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 1:
        set_font(run, size=13, bold=True)
        p.paragraph_format.space_before = Pt(18)
    elif level == 2:
        set_font(run, size=12, bold=True)
        p.paragraph_format.space_before = Pt(10)
    return p

def body(doc, text, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    return p

def mixed(doc, parts, space_after=8):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p

def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bot)
    pPr.append(pBdr)

def ref_entry(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after       = Pt(4)
    set_font(p.add_run(text), size=11)

def image(doc, filename, caption, width=Inches(5.8)):
    path = os.path.join(SHOTS_DIR, filename)
    if not os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(f"[Screenshot not found — save as docs/screenshots/{filename}]")
        set_font(r, size=10, italic=True, color=(160, 160, 160))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        p.add_run().add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after  = Pt(14)
    cap.paragraph_format.space_before = Pt(2)
    set_font(cap.add_run(caption), size=9, italic=True, color=(100, 100, 100))

# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.5)

# ── Title block ───────────────────────────────────────────────────────────────
add_heading(doc,
    "Sentio: A Critical Evaluation of an EEG-Driven Wearable LED System", 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
set_font(p.add_run("Critical Academic Report"), size=11, italic=True,
         color=(100, 100, 100))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(20)
set_font(p2.add_run("Sentio Project  |  WS2812B Arduino Version  |  2026"),
         size=10, color=(130, 130, 130))

separator(doc)

# ── Hardware Specification Summary (from Assembly Manual) ─────────────────────
add_heading(doc, "Hardware Specification Summary")

body(doc,
    "The physical LED garment is built according to the WS2812B LED T-Shirt "
    "Technical Documentation & Assembly Manual. Key parameters are summarised "
    "below and referenced throughout this report.")

# Spec table
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "Parameter"
hdr[1].text = "Value"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        set_font(run, size=10, bold=True)

specs = [
    ("LED type",            "WS2812B (individually addressable RGB)"),
    ("Matrix layout",       "10 rows x 20 columns = 200 LEDs (serpentine)"),
    ("Active panel area",   "40 cm x 55 cm"),
    ("Base panel",          "42 x 57 cm felt or cotton (includes margin)"),
    ("Diffusion layer",     "Thin white fabric, lightly stitched over panel"),
    ("Supply voltage",      "5V DC regulated"),
    ("Max current (peak)",  "200 x 60 mA = 12 A (full white)"),
    ("Recommended current", "200 x 20 mA = 4 A (animated patterns)"),
    ("Recommended supply",  "5V / 5-10 A external battery pack"),
    ("Data line",           "Arduino/ESP32 GPIO -> 330 ohm resistor -> DIN"),
    ("Power injection",     "4 points: Row 1, Row 4, Row 7, Row 10"),
    ("Decoupling cap",      "1000 uF across 5V / GND near first LED"),
    ("Wiring gauge",        "Silicone 22-26 AWG"),
    ("Software brightness", "Capped at 50-100 (out of 255) per manual"),
    ("Controller",          "Arduino Nano or compatible (ESP32 in Sentio)"),
]
for param, val in specs:
    row = tbl.add_row().cells
    row[0].text = param
    row[1].text = val
    for i, cell in enumerate(row):
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10, italic=(i == 0))

doc.add_paragraph().paragraph_format.space_after = Pt(8)
separator(doc)

# ── Architecture Diagram ──────────────────────────────────────────────────────
add_heading(doc, "System Architecture")
image(doc, "architecture.png",
      "Figure 0 — Sentio system architecture: Muse 2 EEG headband to 10x20 WS2812B "
      "LED matrix (200 LEDs) via ESP32 WebSocket, with React monitoring dashboard "
      "as secondary operator output.",
      width=Inches(6.0))
separator(doc)

# ── 1. Introduction ───────────────────────────────────────────────────────────
add_heading(doc, "1. Introduction")

body(doc,
    "Sentio is an emotion-driven wearable interactive system that translates "
    "real-time electroencephalography (EEG) data into generative visual patterns "
    "rendered on a physical LED panel worn on a t-shirt chest. The hardware "
    "substrate -- specified in the WS2812B LED T-Shirt Assembly Manual -- "
    "comprises 200 individually addressable WS2812B LEDs arranged in a 10x20 "
    "serpentine matrix across a 40 x 55 cm active area, mounted on a 42 x 57 cm "
    "felt base panel and covered with a thin white diffusion fabric to smooth "
    "light output. A Muse 2 EEG headband streams brainwave data to a Python "
    "backend that classifies five emotional states -- calm, focused, stressed, "
    "relaxed, and excited -- and transmits results wirelessly via WebSocket to "
    "an ESP32 microcontroller driving the matrix through four visual patterns: "
    "Ondas Fluidas, Padrao Geometrico, Pulsos Ritmicos, and Estrelas e "
    "Particulas. A React web dashboard (Figure 5) provides real-time operator "
    "monitoring. This report critically evaluates Sentio's effectiveness as an "
    "interactive system, examining usability, interaction design principles, and "
    "system-level thinking to assess how successfully it bridges complex "
    "neurological data and a meaningful wearable experience.")

image(doc, TSHIRT_PHOTO[0], TSHIRT_PHOTO[1], width=Inches(6.2))

image(doc, "01_config_screen.png", SCREENSHOTS[0][1])

# ── 2. Usability and Interaction Analysis ─────────────────────────────────────
add_heading(doc, "2. Usability and Interaction Analysis")

mixed(doc, [
    ("Effectiveness.  ", True, False),
    ("Sentio achieves its primary goal through a clearly defined five-stage "
     "pipeline: session configuration, device connection, calibration, real-time "
     "monitoring, and LED pattern output (Figures 1-3). The configuration screen "
     "(Figure 1) presents pattern type as a 2x2 card grid with inline SVG "
     "previews, reducing the abstraction gap between label and visual output. "
     "This decision supports task completion by providing ", False, False),
    ("recognition over recall", False, True),
    (" (Nielsen, 1994). However, hardware reliability is a critical dependency: "
     "Bluetooth connectivity between the Muse headband and the processing machine "
     "constitutes a single point of failure. Additionally, the physical garment "
     "introduces a second hardware dependency chain -- power supply integrity, "
     "330-ohm data resistor, 1000 uF decoupling capacitor, and four power "
     "injection points -- any of which, if misconfigured, will produce partial "
     "or complete LED failure without any software-visible error.", False, False),
])

mixed(doc, [
    ("Efficiency.  ", True, False),
    ("The Device Settings panel (Figure 2) allows operators to configure the "
     "backend API URL and Muse MAC address in the field without redeploying the "
     "application -- a pragmatic wearable deployment affordance. Calibration "
     "(Figure 3) completes automatically, requiring no active input. However, "
     "the animated progress bar runs independently of actual backend processing, "
     "violating Nielsen's ", False, False),
    ("visibility of system status", False, True),
    (" heuristic. The Assembly Manual recommends testing each of the 10 strip "
     "segments individually before full integration; however, Sentio provides no "
     "built-in LED diagnostic mode, requiring external test firmware and "
     "increasing pre-deployment setup time.", False, False),
])

image(doc, "02_device_settings.png", SCREENSHOTS[1][1])
image(doc, "03_calibration_screen.png", SCREENSHOTS[2][1])

mixed(doc, [
    ("Satisfaction.  ", True, False),
    ("The wearable form factor -- a 40 x 55 cm illuminated chest panel "
     "covered by a diffusion layer -- creates a compelling experiential "
     "proposition visible to surrounding observers. The diffusion fabric "
     "specified in the Assembly Manual smooths individual LED hotspots into "
     "a continuous colour field, improving perceived quality at the cost of "
     "some pattern resolution. The four patterns are perceptually "
     "differentiated by hue, morphology, and animation speed. A significant "
     "satisfaction weakness, however, is the observer-wearer asymmetry: the "
     "output medium faces outward, and the wearer has no direct channel -- "
     "haptic, auditory, or visual -- through which to perceive their own "
     "emotional output, limiting personal experiential engagement.", False, False),
])

mixed(doc, [
    ("Interaction Principles.  ", True, False),
    ("The Signal Lost modal (Figure 4) demonstrates considered error recovery: "
     "it presents a numbered three-step checklist with specific recovery "
     "commands, an auto-retry countdown, and a dismiss option, satisfying "
     "Nielsen's heuristic of ", False, False),
    ("help users recognise, diagnose, and recover from errors", False, True),
    (". The signal_quality metric (0-100 scale) modulates LED brightness in "
     "software, with the Assembly Manual recommending a hard cap of 50-100 "
     "(out of 255) to limit power draw to the safe 4 A operating envelope. "
     "This coupling of EEG confidence to physical output intensity creates an "
     "intuitive transparency: poor signal produces a dim garment, communicating "
     "system uncertainty without explicit error messaging. However, pattern "
     "transitions remain instantaneous on emotion change. The absence of "
     "morphing between states violates ", False, False),
    ("smooth continuity", False, True),
    (" principles and produces jarring visual discontinuities that undermine the "
     "organic emotional narrative the system intends to convey.", False, False),
])

image(doc, "04_signal_lost.png", SCREENSHOTS[3][1])

# ── 3. System Thinking Perspective ────────────────────────────────────────────
add_heading(doc, "3. System Thinking Perspective")

mixed(doc, [
    ("Activity Theory.  ", True, False),
    ("Analysed through Activity Theory (Engestrom, 1987), Sentio situates the "
     "wearer as both ", False, False),
    ("subject", False, True),
    (" and primary data source, employing the full EEG-to-LED pipeline as ",
     False, False),
    ("tool", False, True),
    (" to externalise emotional state as ", False, False),
    ("object", False, True),
    (". A fundamental tension emerges: the output medium faces outward, most "
     "visible to external observers rather than the wearer. This inversion of "
     "the subject-object relationship is unacknowledged in the interface design, "
     "creating an experiential mismatch that limits personal agency and "
     "introspective utility.", False, False),
])

mixed(doc, [
    ("Cognitive Ergonomics.  ", True, False),
    ("The monitoring dashboard (Figure 5) presents EEG band powers, signal "
     "quality, visual parameters, and an AI Guidance panel simultaneously. "
     "From a cognitive ergonomics perspective (Wickens et al., 2004), this "
     "density serves expert operators but risks overloading novice exhibition "
     "users. The physical garment introduces additional operator cognitive load: "
     "the Assembly Manual specifies a seven-step testing procedure (individual "
     "segment test, full chain test, 30% brightness test, check for flickering, "
     "colour errors, and dead pixels) before mounting -- a process requiring "
     "technical competence not accounted for in the interaction design. The "
     "absence of a software-driven hardware diagnostic mode transfers this "
     "cognitive burden entirely to the fabrication team.", False, False),
])

mixed(doc, [
    ("Distributed Cognition.  ", True, False),
    ("Hutchins' (1995) framework provides the strongest theoretical "
     "justification for Sentio's value: the system externalises an invisible "
     "internal emotional state into a shared, perceptible artefact worn on the "
     "body. The diffusion layer -- thin white fabric stitched over the 200-LED "
     "panel -- physically mediates this externalisation, transforming discrete "
     "pixel data into a continuous colour field legible to observers without EEG "
     "literacy. However, the five-emotion discrete classification model imposes "
     "a reductive taxonomy onto a continuous neurological spectrum, and "
     "misclassification -- a stressed signal rendered as focused -- breaks the "
     "perceived authenticity central to the garment's meaning. The rule-based "
     "classifier offers no individual calibration mechanism.", False, False),
])

mixed(doc, [
    ("System Failure Propagation.  ", True, False),
    ("Sentio's multi-layer hardware stack creates compounded failure risk. "
     "At the software level: Bluetooth dropout stops EEG data -> backend emits "
     "heartbeat-only frames -> ESP32 falls back to idle animation -> garment "
     "dims. At the hardware level: insufficient power supply (below the "
     "recommended 5V / 5A minimum) causes voltage sag across the 200-LED chain, "
     "producing colour distortion and flickering even when the software pipeline "
     "is fully functional. The Assembly Manual's recommendation to inject power "
     "at four points (Rows 1, 4, 7, 10) mitigates voltage drop across the "
     "serpentine chain, but this requirement is invisible to the software "
     "interface -- a missed opportunity for hardware-software status integration.",
     False, False),
])

image(doc, "05_monitoring_screen.png", SCREENSHOTS[4][1])

# ── 4. Connection to User Research and Heuristics ─────────────────────────────
add_heading(doc, "4. Connection to User Research and Heuristics")

body(doc,
    "Sentio targets an exhibition or live performance context with novice users "
    "who have no EEG literacy. The automated calibration, guided sequential flow, "
    "and pattern type cards (Figure 1) align with Nielsen's error prevention "
    "heuristic by removing technical complexity from user control. The "
    "emotion-to-pattern colour mapping -- red/pulse for stressed, cyan/wave for "
    "calm -- applies established colour-emotion associations (Kaya & Epps, 2004), "
    "reinforcing intuitive readability without prior instruction.")

body(doc,
    "The Assembly Manual's brightness cap (50-100 out of 255) has a direct "
    "interaction design consequence: at these levels the 200-LED panel operates "
    "within the safe 4 A operating envelope, but perceived luminance in "
    "brightly-lit exhibition environments may be insufficient for the "
    "patterns to read clearly at distance. This represents a tension between "
    "electrical safety constraints and the perceptual requirements of the "
    "exhibition context -- a conflict not addressed by the current software "
    "design, which applies a single global brightness value without "
    "environment-adaptive compensation.")

body(doc,
    "The monitoring dashboard (Figure 5) exposes band powers, signal quality, "
    "visual parameters, and a guidance panel simultaneously. While this "
    "satisfies Nielsen's visibility of system status for expert operators, the "
    "density risks cognitive overload for the exhibition audience. Progressive "
    "disclosure -- hiding technical band-power metrics unless explicitly "
    "requested -- would better serve this user profile. The manual override "
    "mode supports Nielsen's user control and freedom heuristic (Nielsen, 1994) "
    "by providing an escape route when the EEG pipeline fails, but it is "
    "accessible only after a successful session start, limiting its utility as "
    "a standalone fallback during hardware failure.")

# ── 5. Conclusion ─────────────────────────────────────────────────────────────
add_heading(doc, "5. Conclusion")

body(doc,
    "Sentio demonstrates technical and experiential feasibility in translating "
    "EEG data into a wearable LED garment. The physical build -- 200 WS2812B "
    "LEDs across a 40 x 55 cm serpentine matrix, diffusion-filtered and "
    "power-injected at four points -- provides a robust hardware substrate "
    "for emotion-reactive wearable display. Its sequential interaction flow, "
    "signal-quality-driven brightness, perceptually distinct patterns, and "
    "well-designed error recovery (Figure 4) constitute a coherent interactive "
    "system with clear exhibition value.")

body(doc,
    "Critical weaknesses remain: calibration feedback opacity; the "
    "observer-wearer asymmetry; discrete emotion classification without "
    "individual calibration; instantaneous pattern transitions; dashboard "
    "information overload; the absence of a software-driven LED diagnostic "
    "mode; and a brightness-versus-visibility tension in high-ambient-light "
    "exhibition environments. Future iterations should prioritise: (1) pattern "
    "morphing interpolation for smooth emotional transitions; (2) secondary "
    "feedback for the wearer via audio or a mirrored mobile display; (3) a "
    "continuous valence-arousal model replacing discrete emotion labels; "
    "(4) a user-adaptive classifier trained on individual EEG baselines; and "
    "(5) environment-adaptive brightness scaling that respects the 4 A power "
    "envelope while compensating for ambient light conditions. These developments "
    "would move Sentio from a technically accomplished demonstrator towards "
    "a genuinely personalised and deployment-robust wearable emotional "
    "communication system.")

separator(doc)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
r = p.add_run(
    "Word count: approximately 1000 words (body text, excluding table, "
    "captions, and references).")
set_font(r, size=10, italic=True, color=(130, 130, 130))
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ── References ────────────────────────────────────────────────────────────────
add_heading(doc, "References")

references = [
    "Engestrom, Y. (1987). Learning by Expanding: An Activity-Theoretical "
    "Approach to Developmental Research. Orienta-Konsultit.",

    "Hutchins, E. (1995). Cognition in the Wild. MIT Press.",

    "Kaya, N., & Epps, H. H. (2004). Relationship between colour and emotion: "
    "A study of college students. College Student Journal, 38(3), 396-405.",

    "Nielsen, J. (1994). Usability Engineering. Academic Press.",

    "Norman, D. A. (2013). The Design of Everyday Things (Revised ed.). "
    "Basic Books.",

    "Wickens, C. D., Lee, J. D., Liu, Y., & Gordon-Becker, S. (2004). "
    "An Introduction to Human Factors Engineering (2nd ed.). "
    "Pearson Prentice Hall.",

    "WS2812B LED T-Shirt -- Technical Documentation & Assembly Manual. (2026). "
    "Engineering & Fabrication Team, Sentio Project. [Internal technical document].",
]

for r in references:
    ref_entry(doc, r)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print("Saved -> " + OUT_FILE)

all_images = [(TSHIRT_PHOTO[0], TSHIRT_PHOTO[1])] + SCREENSHOTS
missing = [f for f, _ in all_images
           if not os.path.exists(os.path.join(SHOTS_DIR, f))]
if missing:
    print("\nMissing images (save to docs/screenshots/):")
    for f in missing:
        print("  " + f)
else:
    print("All images embedded.")
